"""
Final Resume Generator - Streamlined version with GitHub hyperlinks
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def create_resume():
    """Create streamlined ATS-optimized resume"""
    doc = Document()

    # Set narrow margins for single page
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # NAME
    name = doc.add_paragraph()
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = name.add_run('NATASHA XAVIER')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'Calibri'

    # Contact line 1
    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact.add_run('London, United Kingdom | ').font.size = Pt(10)
    add_hyperlink(contact, 'natashaxav.github.io', 'https://natashaxav.github.io')
    contact.add_run(' | ').font.size = Pt(10)
    add_hyperlink(contact, 'github.com/natashaxav', 'https://github.com/natashaxav')

    # Contact line 2
    contact2 = doc.add_paragraph()
    contact2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_hyperlink(contact2, 'natashaxavier512@gmail.com', 'mailto:natashaxavier512@gmail.com')
    contact2.add_run(' | ').font.size = Pt(10)
    add_hyperlink(contact2, 'linkedin.com/in/natashaxavier-marketingspecialist', 'https://linkedin.com/in/natashaxavier-marketingspecialist')
    contact2.paragraph_format.space_after = Pt(6)

    # Separator
    separator = doc.add_paragraph('_' * 90)
    separator.paragraph_format.space_before = Pt(2)
    separator.paragraph_format.space_after = Pt(6)
    separator.runs[0].font.size = Pt(8)
    separator.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # EDUCATION
    edu_heading = doc.add_paragraph()
    edu_run = edu_heading.add_run('EDUCATION')
    edu_run.font.size = Pt(12)
    edu_run.font.bold = True
    edu_run.font.name = 'Calibri'
    edu_heading.paragraph_format.space_before = Pt(4)
    edu_heading.paragraph_format.space_after = Pt(4)

    # Master's
    masters = doc.add_paragraph()
    masters_run = masters.add_run('Master of Science, Marketing Management')
    masters_run.font.size = Pt(10.5)
    masters_run.font.bold = True
    masters_run.font.name = 'Calibri'
    masters.paragraph_format.space_after = Pt(1)

    masters_details = doc.add_paragraph('University of Westminster, London, UK | December 2023')
    masters_details.runs[0].font.size = Pt(10)
    masters_details.runs[0].font.name = 'Calibri'
    masters_details.paragraph_format.space_after = Pt(1)

    masters_coursework = doc.add_paragraph('Relevant Coursework: Marketing Research Data Analytics, Multiplatform Marketing, Consumer Psychology, Digital Strategy')
    masters_coursework.runs[0].font.size = Pt(9.5)
    masters_coursework.runs[0].font.italic = True
    masters_coursework.runs[0].font.name = 'Calibri'
    masters_coursework.paragraph_format.space_after = Pt(6)

    # Bachelor's
    bachelors = doc.add_paragraph()
    bachelors_run = bachelors.add_run('Bachelor of Business Administration, International Business')
    bachelors_run.font.size = Pt(10.5)
    bachelors_run.font.bold = True
    bachelors_run.font.name = 'Calibri'
    bachelors.paragraph_format.space_after = Pt(1)

    bachelors_details = doc.add_paragraph("St. Joseph's College of Commerce, Bangalore, IN | June 2020")
    bachelors_details.runs[0].font.size = Pt(10)
    bachelors_details.runs[0].font.name = 'Calibri'
    bachelors_details.paragraph_format.space_after = Pt(1)

    bachelors_coursework = doc.add_paragraph('Relevant Coursework: Marketing Management, Project Management, International Business, Strategic Planning')
    bachelors_coursework.runs[0].font.size = Pt(9.5)
    bachelors_coursework.runs[0].font.italic = True
    bachelors_coursework.runs[0].font.name = 'Calibri'
    bachelors_coursework.paragraph_format.space_after = Pt(6)

    # Separator
    separator2 = doc.add_paragraph('_' * 90)
    separator2.paragraph_format.space_before = Pt(2)
    separator2.paragraph_format.space_after = Pt(6)
    separator2.runs[0].font.size = Pt(8)
    separator2.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # TECHNICAL SKILLS
    skills_heading = doc.add_paragraph()
    skills_run = skills_heading.add_run('TECHNICAL SKILLS')
    skills_run.font.size = Pt(12)
    skills_run.font.bold = True
    skills_run.font.name = 'Calibri'
    skills_heading.paragraph_format.space_before = Pt(4)
    skills_heading.paragraph_format.space_after = Pt(4)

    skills_list = [
        'AI & Automation: CrewAI, LangChain, Gemini AI, Claude AI, automated reporting, sentiment analysis',
        'Digital Marketing: SEMrush, Google Analytics, Hootsuite, Canva, campaign management, SEO optimization',
        'Data & Analytics: Tableau, Excel, Google Data Studio, RFM analysis, customer segmentation',
        'HR & CRM Tools: Workday, ATS platforms, Zapier, CRM software, HRIS systems'
    ]

    for skill in skills_list:
        p = doc.add_paragraph(skill, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = 'Calibri'

    # Separator
    separator3 = doc.add_paragraph('_' * 90)
    separator3.paragraph_format.space_before = Pt(4)
    separator3.paragraph_format.space_after = Pt(6)
    separator3.runs[0].font.size = Pt(8)
    separator3.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # PROFESSIONAL EXPERIENCE
    exp_heading = doc.add_paragraph()
    exp_run = exp_heading.add_run('PROFESSIONAL EXPERIENCE')
    exp_run.font.size = Pt(12)
    exp_run.font.bold = True
    exp_run.font.name = 'Calibri'
    exp_heading.paragraph_format.space_before = Pt(4)
    exp_heading.paragraph_format.space_after = Pt(4)

    # Job 1
    job1_title = doc.add_paragraph()
    job1_title_run = job1_title.add_run('Social Media Marketing Intern ')
    job1_title_run.font.size = Pt(10.5)
    job1_title_run.font.bold = True
    job1_title_run.font.name = 'Calibri'
    job1_company = job1_title.add_run('| Pictures in Motion, London, UK')
    job1_company.font.size = Pt(10.5)
    job1_company.font.name = 'Calibri'
    job1_title.paragraph_format.space_after = Pt(1)

    job1_date = doc.add_paragraph('April 2024 – Present')
    job1_date.runs[0].font.size = Pt(10)
    job1_date.runs[0].font.italic = True
    job1_date.runs[0].font.name = 'Calibri'
    job1_date.paragraph_format.space_after = Pt(2)

    job1_bullets = [
        'Boosted engagement 25% and audience 20% through strategic content planning and performance tracking',
        'Delivered weekly analytics reports identifying optimization opportunities for marketing campaigns',
        'Coordinated with teams to maximize press visibility and audience reach for content launches'
    ]

    for bullet in job1_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = 'Calibri'

    # Job 2
    job2_title = doc.add_paragraph()
    job2_title.paragraph_format.space_before = Pt(4)
    job2_title_run = job2_title.add_run('Marketing Analyst ')
    job2_title_run.font.size = Pt(10.5)
    job2_title_run.font.bold = True
    job2_title_run.font.name = 'Calibri'
    job2_company = job2_title.add_run('| Grey Apple Advertising, Remote')
    job2_company.font.size = Pt(10.5)
    job2_company.font.name = 'Calibri'
    job2_title.paragraph_format.space_after = Pt(1)

    job2_date = doc.add_paragraph('February 2022 – February 2023')
    job2_date.runs[0].font.size = Pt(10)
    job2_date.runs[0].font.italic = True
    job2_date.runs[0].font.name = 'Calibri'
    job2_date.paragraph_format.space_after = Pt(2)

    job2_bullets = [
        'Increased organic lead generation 12% through strategic keyword research and content optimization',
        'Delivered 150+ qualified leads per event through market analysis and strategic event planning',
        'Enhanced brand engagement 25% through targeted audience campaigns and digital strategy'
    ]

    for bullet in job2_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = 'Calibri'

    # Job 3
    job3_title = doc.add_paragraph()
    job3_title.paragraph_format.space_before = Pt(4)
    job3_title_run = job3_title.add_run('Human Resources Analyst ')
    job3_title_run.font.size = Pt(10.5)
    job3_title_run.font.bold = True
    job3_title_run.font.name = 'Calibri'
    job3_company = job3_title.add_run('| Walmart Global Tech, Bangalore, IN')
    job3_company.font.size = Pt(10.5)
    job3_company.font.name = 'Calibri'
    job3_title.paragraph_format.space_after = Pt(1)

    job3_date = doc.add_paragraph('October 2020 – February 2022')
    job3_date.runs[0].font.size = Pt(10)
    job3_date.runs[0].font.italic = True
    job3_date.runs[0].font.name = 'Calibri'
    job3_date.paragraph_format.space_after = Pt(2)

    job3_bullets = [
        'Streamlined recruitment process by implementing automation system, reducing hiring time 23%',
        'Designed compensation framework achieving 95% satisfaction and 10% lower staff turnover',
        'Improved team engagement 15% through strategic feedback initiatives and communication programs'
    ]

    for bullet in job3_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = 'Calibri'

    # Separator
    separator4 = doc.add_paragraph('_' * 90)
    separator4.paragraph_format.space_before = Pt(4)
    separator4.paragraph_format.space_after = Pt(6)
    separator4.runs[0].font.size = Pt(8)
    separator4.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # KEY PROJECTS
    projects_heading = doc.add_paragraph()
    projects_run = projects_heading.add_run('KEY PROJECTS')
    projects_run.font.size = Pt(12)
    projects_run.font.bold = True
    projects_run.font.name = 'Calibri'
    projects_heading.paragraph_format.space_before = Pt(4)
    projects_heading.paragraph_format.space_after = Pt(4)

    # Projects with hyperlinks
    projects = [
        {
            'name': 'Social Media Sentiment Analysis',
            'url': 'https://github.com/natashaxav/social-media-sentiment-analysis',
            'desc': 'Automated brand monitoring system analyzing audience sentiment patterns using Gemini AI, achieving 87% accuracy and £5,000+ annual cost savings'
        },
        {
            'name': 'Customer Segmentation - RFM Analysis',
            'url': 'https://github.com/natashaxav/customer-segmentation-rfm',
            'desc': 'Analyzed 973 customers across £685K revenue using behavioral clustering to enable personalized campaign strategies and improve retention'
        },
        {
            'name': 'AI Competitor Intelligence Analyzer',
            'url': 'https://github.com/natashaxav/ai-competitor-analyzer',
            'desc': 'Deployed AI agents via CrewAI and Gemini to automate market intelligence gathering, delivering 20x ROI versus traditional research services'
        },
        {
            'name': 'Marketing Content Generator',
            'url': 'https://github.com/natashaxav/marketing-content-generator',
            'desc': 'Built AI-powered content system using Gemini AI increasing brand-aligned output 4x and saving 15 hours weekly'
        },
        {
            'name': 'HR Resume Screener',
            'url': 'https://github.com/natashaxav/hr-resume-screener',
            'desc': 'Automated candidate evaluation using AI-powered analysis, reducing screening time 90% and accelerating hiring by 40%'
        },
        {
            'name': 'Employee Engagement Analyzer',
            'url': 'https://github.com/natashaxav/employee-engagement-analyzer',
            'desc': 'Implemented automated survey analysis using Gemini AI to identify retention risks, saving 40 hours per survey cycle'
        }
    ]

    for project in projects:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)

        # Project name with hyperlink
        run = p.add_run(project['name'])
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.name = 'Calibri'

        p.add_run(' (')
        add_hyperlink(p, 'GitHub', project['url'])
        p.add_run(')\n')

        # Description
        desc_run = p.add_run(project['desc'])
        desc_run.font.size = Pt(9.5)
        desc_run.font.name = 'Calibri'

    # Save the document
    output_path = '/Users/rj/Claude_code/natash/Natasha_Xavier_Resume_FINAL.docx'
    doc.save(output_path)
    print(f"✓ Professional DOCX created: {output_path}")
    return output_path

if __name__ == '__main__':
    try:
        print("Creating streamlined ATS-optimized resume...")
        docx_path = create_resume()
        print("\n✓ Resume created successfully!")
        print(f"\nDOCX file: {docx_path}")
        print("\nTo convert to PDF:")
        print("Open in Microsoft Word/Google Docs → File → Save As → PDF")

    except Exception as e:
        print(f"Error creating resume: {str(e)}")
        import traceback
        traceback.print_exc()
