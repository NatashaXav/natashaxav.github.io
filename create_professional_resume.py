"""
Professional Resume Generator - Creates ATS-optimized DOCX and PDF
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph
    """
    # This gets access to the document.xml.rels file
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a new run
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set the run's style to Hyperlink
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

    return hyperlink

def create_resume():
    """
    Create professional ATS-optimized resume
    """
    doc = Document()

    # Set narrow margins for single page
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # NAME - Large, bold, centered
    name = doc.add_paragraph()
    name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = name.add_run('NATASHA XAVIER')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = 'Calibri'

    # Contact info - centered, with hyperlinks
    contact = doc.add_paragraph()
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_run = contact.add_run('London, United Kingdom | ')
    contact_run.font.size = Pt(10)
    contact_run.font.name = 'Calibri'

    # Email hyperlink
    add_hyperlink(contact, 'natashaxavier512@gmail.com', 'mailto:natashaxavier512@gmail.com')
    contact.add_run(' | ').font.size = Pt(10)

    # Portfolio hyperlink
    add_hyperlink(contact, 'natashaxav.github.io', 'https://natashaxav.github.io')
    contact.add_run('\n').font.size = Pt(10)

    # Second line of contact - LinkedIn and GitHub
    contact2 = doc.add_paragraph()
    contact2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_hyperlink(contact2, 'LinkedIn', 'https://linkedin.com/in/natashaxavier-marketingspecialist')
    contact2.add_run(' | ').font.size = Pt(10)
    add_hyperlink(contact2, 'GitHub', 'https://github.com/natashaxav')

    # Add small spacing
    contact2.paragraph_format.space_after = Pt(6)

    # Separator line
    separator = doc.add_paragraph('_' * 90)
    separator.paragraph_format.space_before = Pt(2)
    separator.paragraph_format.space_after = Pt(6)
    run = separator.runs[0]
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # EDUCATION SECTION
    edu_heading = doc.add_paragraph()
    edu_run = edu_heading.add_run('EDUCATION')
    edu_run.font.size = Pt(12)
    edu_run.font.bold = True
    edu_run.font.name = 'Calibri'
    edu_heading.paragraph_format.space_before = Pt(4)
    edu_heading.paragraph_format.space_after = Pt(4)

    # Master's degree
    masters = doc.add_paragraph()
    masters_run = masters.add_run('Master of Science, Marketing Management')
    masters_run.font.size = Pt(10.5)
    masters_run.font.bold = True
    masters_run.font.name = 'Calibri'
    masters.paragraph_format.space_after = Pt(1)

    masters_details = doc.add_paragraph('University of Westminster, London, UK | December 2023')
    masters_details.paragraph_format.space_after = Pt(6)
    masters_details.runs[0].font.size = Pt(10)
    masters_details.runs[0].font.name = 'Calibri'

    # Bachelor's degree
    bachelors = doc.add_paragraph()
    bachelors_run = bachelors.add_run('Bachelor of Business Administration, International Business')
    bachelors_run.font.size = Pt(10.5)
    bachelors_run.font.bold = True
    bachelors_run.font.name = 'Calibri'
    bachelors.paragraph_format.space_after = Pt(1)

    bachelors_details = doc.add_paragraph("St. Joseph's College of Commerce, Bangalore, IN | June 2020")
    bachelors_details.paragraph_format.space_after = Pt(6)
    bachelors_details.runs[0].font.size = Pt(10)
    bachelors_details.runs[0].font.name = 'Calibri'

    # Separator
    separator2 = doc.add_paragraph('_' * 90)
    separator2.paragraph_format.space_before = Pt(2)
    separator2.paragraph_format.space_after = Pt(6)
    run = separator2.runs[0]
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # TECHNICAL SKILLS SECTION
    skills_heading = doc.add_paragraph()
    skills_run = skills_heading.add_run('TECHNICAL SKILLS')
    skills_run.font.size = Pt(12)
    skills_run.font.bold = True
    skills_run.font.name = 'Calibri'
    skills_heading.paragraph_format.space_before = Pt(4)
    skills_heading.paragraph_format.space_after = Pt(4)

    skills_list = [
        'Programming & AI: Python (pandas, NumPy, scikit-learn), SQL',
        'NLP & Machine Learning: NLTK, spaCy, Transformers, BERT, Sentiment Analysis, Topic Modeling (LDA, BERTopic), K-Means Clustering',
        'AI Frameworks: CrewAI, LangChain, GPT-4, Gemini, Claude',
        'Data Visualization: Tableau, Chart.js, Plotly, Advanced Excel',
        'Marketing & HR Tech: Google Analytics, SEMrush, Workday, ATS, Zapier'
    ]

    for skill in skills_list:
        p = doc.add_paragraph(skill, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = 'Calibri'

    # Separator
    separator3 = doc.add_paragraph('_' * 90)
    separator3.paragraph_format.space_before = Pt(4)
    separator3.paragraph_format.space_after = Pt(6)
    run = separator3.runs[0]
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # PROFESSIONAL EXPERIENCE SECTION
    exp_heading = doc.add_paragraph()
    exp_run = exp_heading.add_run('PROFESSIONAL EXPERIENCE')
    exp_run.font.size = Pt(12)
    exp_run.font.bold = True
    exp_run.font.name = 'Calibri'
    exp_heading.paragraph_format.space_before = Pt(4)
    exp_heading.paragraph_format.space_after = Pt(4)

    # Job 1: Pictures in Motion
    job1_title = doc.add_paragraph()
    job1_title_run = job1_title.add_run('Social Media Marketing Intern ')
    job1_title_run.font.size = Pt(10.5)
    job1_title_run.font.bold = True
    job1_title_run.font.name = 'Calibri'
    job1_company = job1_title.add_run('| Pictures in Motion, London, UK')
    job1_company.font.size = Pt(10.5)
    job1_company.font.name = 'Calibri'
    job1_title.paragraph_format.space_after = Pt(1)

    job1_date = doc.add_paragraph('April 2024 – Present')
    job1_date.runs[0].font.size = Pt(10)
    job1_date.runs[0].font.italic = True
    job1_date.runs[0].font.name = 'Calibri'
    job1_date.paragraph_format.space_after = Pt(2)

    job1_bullets = [
        'Increased social media engagement by 25% and grew audience by 20% through data-informed content strategy',
        'Analyzed campaign performance using Google Analytics and Hootsuite, delivering actionable weekly reports',
        'Collaborated across teams to amplify press releases and drive engagement for TV show launches'
    ]

    for bullet in job1_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = 'Calibri'

    # Job 2: Grey Apple Advertising
    job2_title = doc.add_paragraph()
    job2_title.paragraph_format.space_before = Pt(4)
    job2_title_run = job2_title.add_run('Marketing Analyst ')
    job2_title_run.font.size = Pt(10.5)
    job2_title_run.font.bold = True
    job2_title_run.font.name = 'Calibri'
    job2_company = job2_title.add_run('| Grey Apple Advertising, Remote')
    job2_company.font.size = Pt(10.5)
    job2_company.font.name = 'Calibri'
    job2_title.paragraph_format.space_after = Pt(1)

    job2_date = doc.add_paragraph('February 2022 – February 2023')
    job2_date.runs[0].font.size = Pt(10)
    job2_date.runs[0].font.italic = True
    job2_date.runs[0].font.name = 'Calibri'
    job2_date.paragraph_format.space_after = Pt(2)

    job2_bullets = [
        'Drove 12% increase in organic leads through SEO strategies, improving rankings for top keywords',
        'Generated 150+ qualified leads per event through three high-impact media branding events',
        'Enhanced customer engagement by 25% through targeted CRM campaigns and brand-aligned content'
    ]

    for bullet in job2_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = 'Calibri'

    # Job 3: Walmart Global Tech
    job3_title = doc.add_paragraph()
    job3_title.paragraph_format.space_before = Pt(4)
    job3_title_run = job3_title.add_run('Human Resources Analyst ')
    job3_title_run.font.size = Pt(10.5)
    job3_title_run.font.bold = True
    job3_title_run.font.name = 'Calibri'
    job3_company = job3_title.add_run('| Walmart Global Tech, Bangalore, IN')
    job3_company.font.size = Pt(10.5)
    job3_company.font.name = 'Calibri'
    job3_title.paragraph_format.space_after = Pt(1)

    job3_date = doc.add_paragraph('October 2020 – February 2022')
    job3_date.runs[0].font.size = Pt(10)
    job3_date.runs[0].font.italic = True
    job3_date.runs[0].font.name = 'Calibri'
    job3_date.paragraph_format.space_after = Pt(2)

    job3_bullets = [
        'Implemented new ATS for tech hub, automating workflows and reducing time-to-hire by 23%',
        'Developed compensation strategies achieving 95% employee satisfaction and 10% lower turnover',
        'Improved employee engagement by 15% through strategic surveys and communication programs'
    ]

    for bullet in job3_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.name = 'Calibri'

    # Separator
    separator4 = doc.add_paragraph('_' * 90)
    separator4.paragraph_format.space_before = Pt(4)
    separator4.paragraph_format.space_after = Pt(6)
    run = separator4.runs[0]
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # KEY PROJECTS SECTION
    projects_heading = doc.add_paragraph()
    projects_run = projects_heading.add_run('KEY PROJECTS')
    projects_run.font.size = Pt(12)
    projects_run.font.bold = True
    projects_run.font.name = 'Calibri'
    projects_heading.paragraph_format.space_before = Pt(4)
    projects_heading.paragraph_format.space_after = Pt(4)

    # Projects with hyperlinks
    projects = [
        {
            'name': 'Social Media Sentiment Analysis',
            'url': 'https://github.com/natashaxav/social-media-sentiment-analysis',
            'desc': 'Built NLP model analyzing 10,000 tweets with 87% accuracy using Python, NLTK, VADER, spaCy, and Transformers'
        },
        {
            'name': 'Customer Segmentation - RFM Analysis',
            'url': 'https://github.com/natashaxav/customer-segmentation-rfm',
            'desc': 'Segmented 973 customers using K-Means clustering and RFM methodology with Python and scikit-learn'
        },
        {
            'name': 'AI Competitor Intelligence Analyzer',
            'url': 'https://github.com/natashaxav/ai-competitor-analyzer',
            'desc': 'Deployed autonomous AI agent using CrewAI and Gemini for competitive intelligence automation'
        },
        {
            'name': 'Marketing Content Generator',
            'url': 'https://github.com/natashaxav/marketing-content-generator',
            'desc': 'Engineered GPT-4 powered system for brand-aligned content creation with prompt engineering'
        },
        {
            'name': 'HR Resume Screener',
            'url': 'https://github.com/natashaxav/hr-resume-screener',
            'desc': 'Developed NLP tool using spaCy and BERT for automated resume parsing and candidate matching'
        },
        {
            'name': 'Employee Engagement Analyzer',
            'url': 'https://github.com/natashaxav/employee-engagement-analyzer',
            'desc': 'Built sentiment analysis system using VADER, TextBlob, and BERTopic for employee survey analysis'
        }
    ]

    for project in projects:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)

        # Add project name as hyperlink in bold
        run = p.add_run(project['name'] + ': ')
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.name = 'Calibri'

        # Add description
        desc_run = p.add_run(project['desc'])
        desc_run.font.size = Pt(10)
        desc_run.font.name = 'Calibri'

        # Add hyperlink at the end
        p.add_run(' (')
        add_hyperlink(p, 'GitHub', project['url'])
        p.add_run(')')

    # Save the document
    output_path = '/Users/rj/Claude_code/natash/Natasha_Xavier_Resume_Professional.docx'
    doc.save(output_path)
    print(f"✓ Professional DOCX created: {output_path}")

    return output_path

if __name__ == '__main__':
    try:
        print("Creating professional ATS-optimized resume...")
        docx_path = create_resume()
        print("\nResume created successfully!")
        print(f"\nDOCX file: {docx_path}")
        print("\nTo convert to PDF:")
        print("1. Open the DOCX file in Microsoft Word or Google Docs")
        print("2. Go to File > Save As > PDF")
        print("3. Or use: pip install docx2pdf && python -c 'from docx2pdf import convert; convert(\"Natasha_Xavier_Resume_Professional.docx\")'")

    except Exception as e:
        print(f"Error creating resume: {str(e)}")
        import traceback
        traceback.print_exc()
